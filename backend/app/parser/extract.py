"""Format readers.

Kept apart from every other package in this repository. Nothing here imports
`app.domain`, `app.infrastructure`, or anything that reaches a database or a
setting, because this code runs in the container that is assumed to fall: it is
the one that reads attacker-supplied bytes with libraries that have a CVE
history. What it cannot import, an exploit in it cannot reach.
"""

from __future__ import annotations

import io

PDF_MEDIA_TYPE = "application/pdf"
DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
TEXT_MEDIA_TYPES = ("text/plain", "text/markdown")

MAX_PAGES = 2000
"""A bound on work, not on document size. A PDF can declare far more pages than
its byte count suggests, and the byte limit upstream does not constrain how long
rendering them takes."""


class UnsupportedMediaType(Exception):
    pass


class ExtractionFailed(Exception):
    pass


def extract(media_type: str, data: bytes) -> str:
    if media_type == PDF_MEDIA_TYPE:
        return _extract_pdf(data)
    if media_type == DOCX_MEDIA_TYPE:
        return _extract_docx(data)
    if media_type in TEXT_MEDIA_TYPES:
        return _extract_text(data)
    raise UnsupportedMediaType(media_type)


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
        pages = reader.pages[:MAX_PAGES]
        return "\n\n".join(page.extract_text() or "" for page in pages)
    except Exception as exc:
        # A bare `except Exception` rather than pypdf's own `PyPdfError`,
        # because pypdf raises plain ValueError, KeyError and friends on
        # malformed input as readily as it raises its own class. A crafted file
        # is the expected case here, not the surprising one, so everything
        # becomes the same refusal and the caller cannot distinguish a crafted
        # file from a merely broken one.
        raise ExtractionFailed(f"pdf: {type(exc).__name__}") from exc


def _extract_docx(data: bytes) -> str:
    import docx

    try:
        document = docx.Document(io.BytesIO(data))
        blocks = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                blocks.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(b for b in blocks if b.strip())
    except Exception as exc:
        raise ExtractionFailed(f"docx: {type(exc).__name__}") from exc


def _extract_text(data: bytes) -> str:
    """`errors="replace"` rather than a refusal: a text file in an unexpected
    encoding is an ordinary event, and mangling a few characters is a better
    outcome than rejecting the upload."""
    return data.decode("utf-8", errors="replace")
